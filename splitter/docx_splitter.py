"""
DOCX 章節偵測器 (Chapter Detector for DOCX)

偵測優先順序：
1. Outline Level (outlineLvl) — 大綱層級（最通用）
2. python-docx Heading 樣式名稱 (Heading 1 / 標題 1)
3. 正則偵測 「第X章」「Chapter X」「# 標題」
"""
import re
import io
from docx import Document
from docx.oxml.ns import qn

# 子章節標題模式：開頭為「數字.數字」代表 2.1、3.2.1 等应排除
_SUB_CHAPTER_RE = re.compile(r"^\d+\.\d")


def _is_sub_chapter(title: str) -> bool:
    return bool(_SUB_CHAPTER_RE.match(title.strip()))

def _get_chapter_prefix(title: str) -> str:
    m = re.match(r"^(第\s*[零一二三四五六七八九十百0-9]+\s*章|Chapter\s+\d+|#{1,2}\s+\d+|\d+\.|Appendix\s+[A-Z0-9]?|Appendices|Solutions|Answers|附錄[A-Z0-9]?|解答)", title, re.IGNORECASE)
    if m:
        return " ".join(m.group(1).split()).lower()
    return title.lower()

def _filter_and_dedup(found: list[tuple]) -> list[tuple]:
    # 尋找 Chapter 1 並捨棄之前的
    first_chap1_idx = -1
    for i, (_, title) in enumerate(found):
        if re.search(r"第\s*[一1]\s*章|Chapter\s+1(?!\d)|^\s*1\.\s+", title, re.IGNORECASE):
            first_chap1_idx = i
            break
            
    if first_chap1_idx != -1:
        found = found[first_chap1_idx:]

    # 依 prefix 去重
    new_found = []
    seen_prefixes = set()
    for item in found:
        prefix = _get_chapter_prefix(item[1])
        if prefix not in seen_prefixes:
            new_found.append(item)
            seen_prefixes.add(prefix)
            
    return new_found


def detect_chapters_docx(file_path: str) -> list[dict]:
    """
    回傳 list of dict:
    [{"title": "第一章 緒論", "start_para": 0, "end_para": 45}, ...]
    段落索引為 0-indexed。
    """
    try:
        doc = Document(file_path)
    except Exception as e:
        err_msg = str(e).lower()
        if "not a zip" in err_msg or "bad zip" in err_msg or "not a word" in err_msg:
            raise ValueError(
                "此檔案可能是舊版 .doc (Office 97) 格式，python-docx 僅支援 .docx。\n"
                "請在 Word 中另存為 .docx 格式後再上傳。"
            ) from e
        raise ValueError(f"無法開啟 DOCX 檔案：{e}") from e

    paragraphs = doc.paragraphs
    total = len(paragraphs)

    if total == 0:
        raise ValueError("此文件無任何段落內容。")

    # ── 方法 1：Outline Level（大綱層級）──────────────────────────
    chapters = _detect_by_outline_level(paragraphs, total)
    if chapters:
        return chapters

    # ── 方法 2：Heading 樣式名稱 ──────────────────────────────────
    chapters = _detect_by_heading(paragraphs, total)
    if chapters:
        return chapters

    # ── 方法 3：正則掃描段落 ──────────────────────────────────────
    chapters = _detect_by_regex(paragraphs, total)
    if chapters:
        return chapters

    raise ValueError(
        "無法偵測章節結構。\n"
        "請確認文件使用「標題1」樣式，或章節標題符合「第X章」/「Chapter X」/「# 標題」格式。"
    )


def _get_outline_level(para) -> int | None:
    """從 XML 的 outlineLvl 取得大綱層級，None 表示無設定。"""
    pPr = para._element.find(qn("w:pPr"))
    if pPr is not None:
        outlineLvl = pPr.find(qn("w:outlineLvl"))
        if outlineLvl is not None:
            val = outlineLvl.get(qn("w:val"))
            if val is not None:
                try:
                    return int(val)
                except ValueError:
                    pass
    return None


def _detect_by_outline_level(paragraphs, total: int) -> list[dict]:
    """使用 outlineLvl XML 屬性偵測大綱層級 0（等同 Heading 1）。"""
    found = []
    seen_titles = set()
    for i, para in enumerate(paragraphs):
        level = _get_outline_level(para)
        if level == 0 and para.text.strip():
            t = para.text.strip()
            if not _is_sub_chapter(t) and t not in seen_titles:
                found.append((i, t))
                seen_titles.add(t)

    if not found:
        return []

    return _build_chapters(found, total)


def _detect_by_heading(paragraphs, total: int) -> list[dict]:
    """使用樣式名稱比對 Heading 1。"""
    HEADING_STYLES = {"heading 1", "標題 1", "heading1", "titre 1", "überschrift 1", "見出し 1"}
    found = []
    seen_titles = set()
    for i, para in enumerate(paragraphs):
        style_name = para.style.name.lower() if para.style else ""
        if style_name in HEADING_STYLES and para.text.strip():
            t = para.text.strip()
            if not _is_sub_chapter(t) and t not in seen_titles:
                found.append((i, t))
                seen_titles.add(t)

    if not found:
        return []

    return _build_chapters(found, total)


def _detect_by_regex(paragraphs, total: int) -> list[dict]:
    """以正則掃描段落文字。"""
    PATTERNS = [
        re.compile(r"^第\s*[零一二三四五六七八九十百0-9]+\s*章\s*.*"),
        re.compile(r"^Chapter\s+\d+(?!\.\d)\s*.*", re.IGNORECASE),
        re.compile(r"^(Appendix\s+[A-Z0-9]?|Appendices|附錄[A-Z0-9]?|解答)\b\s*.*", re.IGNORECASE),
        re.compile(r"^(Solutions|Answers)(?:\s+(to|for|and)\b.*|\s*[:：].*|\s*)$", re.IGNORECASE),
        re.compile(r"^#{1,2}\s+.*"),
    ]
    found = []
    seen_titles = set()
    for i, para in enumerate(paragraphs):
        text = para.text.strip()
        if not text or _is_sub_chapter(text):
            continue
        if text in seen_titles:
            continue
        for p in PATTERNS:
            if p.match(text):
                found.append((i, text))
                seen_titles.add(text)
                break

    if not found:
        return []

    return _build_chapters(found, total)


def _build_chapters(found: list[tuple], total: int) -> list[dict]:
    """從 (para_index, title) 列表建構章節 list。"""
    found = _filter_and_dedup(found)
    chapters = []
    for idx, (para_idx, title) in enumerate(found):
        if idx + 1 < len(found):
            end_para = found[idx + 1][0] - 1
        else:
            end_para = total - 1
        chapters.append({
            "title": title,
            "start_para": para_idx,
            "end_para": end_para,
        })
    return chapters


def extract_chapters_docx(file_path: str, chapters: list[dict], selected_titles: list[str]) -> bytes:
    """
    提取選定章節並合併為新的 DOCX，回傳 bytes。
    使用 XML 層級複製段落以保留完整格式。
    """
    from copy import deepcopy
    from docx.oxml import OxmlElement

    source_doc = Document(file_path)
    paragraphs = source_doc.paragraphs

    new_doc = Document()
    # 清除預設空段落
    for para in new_doc.paragraphs:
        p = para._element
        p.getparent().remove(p)

    chapter_map = {c["title"]: c for c in chapters}
    first = True

    for title in selected_titles:
        if title not in chapter_map:
            continue

        ch = chapter_map[title]
        # 若不是第一章，加分頁
        if not first:
            page_break_para = OxmlElement("w:p")
            pPr = OxmlElement("w:pPr")
            page_break_para.append(pPr)
            run = OxmlElement("w:r")
            br = OxmlElement("w:br")
            br.set(qn("w:type"), "page")
            run.append(br)
            page_break_para.append(run)
            new_doc.element.body.append(page_break_para)

        first = False
        for i in range(ch["start_para"], ch["end_para"] + 1):
            src_para = paragraphs[i]
            # 使用 XML 深拷貝保留完整格式
            new_para_elem = deepcopy(src_para._element)
            new_doc.element.body.append(new_para_elem)

    buf = io.BytesIO()
    new_doc.save(buf)
    return buf.getvalue()
